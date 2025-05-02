import os
import json
import streamlit as st
from groq import Groq
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
from io import BytesIO

# Load API Key
working_dir = os.path.dirname(os.path.abspath(__file__))
config_data = json.load(open(f"{working_dir}/config.json"))
GROQ_API_KEY = config_data["GROQ_API_KEY"]
os.environ["GROQ_API_KEY"] = GROQ_API_KEY
client = Groq(api_key=GROQ_API_KEY)

# Streamlit UI Setup
st.set_page_config(page_title="Case Study Builder", page_icon="📚", layout="centered")
st.markdown("""
<style>
    html, body, [class*="css"]  {
        background-color: #f5faff;
        font-family: 'Segoe UI', sans-serif;
    }
    .stChatMessage.user {
        background-color: #fff3e6;
        border-left: 4px solid #ffa94d;
        padding: 10px;
        border-radius: 10px;
        margin-bottom: 10px;
    }
    .stChatMessage.assistant {
        background-color: #e7f5ff;
        border-left: 4px solid #339af0;
        padding: 10px;
        border-radius: 10px;
        margin-bottom: 10px;
    }
    .stDownloadButton button {
        border-radius: 8px !important;
        background-color: #339af0 !important;
        color: white !important;
        font-weight: 500;
        padding: 6px 12px !important;
    }
    h1 {
        color: #1c7ed6;
        font-size: 2rem !important;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: center;'>📚 Case Study Builder</h1>", unsafe_allow_html=True)
st.markdown("Upload documents or enter a prompt to auto-generate structured case studies for defense training programs.")

with st.expander("📂 Upload Case Study Documents", expanded=True):
    uploaded_files = st.file_uploader("Supported formats: PDF, Word, or Text files", accept_multiple_files=True, type=['pdf', 'docx', 'txt'])

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

def extract_text_from_file(file):
    text = ""
    try:
        if file.type == "application/pdf":
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_file = Document(file)
            for para in docx_file.paragraphs:
                text += para.text + "\n"
        elif file.type.startswith("text/"):
            text = file.read().decode("utf-8")
    except Exception as e:
        st.warning(f"Couldn't read {file.name}: {e}")
    return text

def generate_docx_buffer(content):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Generated Case Study", level=1)
    for section in content.split("\n\n"):
        doc.add_paragraph(section.strip())
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_buffer(content):
    buffer = BytesIO()
    doc = fitz.open()
    page = doc.new_page()
    y = 50
    for paragraph in content.strip().split("\n\n"):
        page.insert_text((50, y), paragraph.strip(), fontsize=11)
        y += 30
        if y > 750:
            page = doc.new_page()
            y = 50
    doc.save(buffer)
    doc.close()
    buffer.seek(0)
    return buffer

system_prompt = ( """
        You are a professional case study generator for a Program Manager leadership course.
        Use the uploaded content to generate a compelling case that follows internationally accepted academic case writing guidelines, combining AMRC, CAIS, and Penn State standards.
        Structure the output using the following six sections:
        1. Opening Paragraph  
        - Clearly state the date, location, and organization.  
        - Identify the main decision-maker by full name and role.  
        - Describe the event or issue trigger that brought the situation to light.  
        - Introduce the core problem or opportunity without revealing the outcome.  
        - Keep it concise and written in past tense, using a neutral and journalistic tone.
        2. Background  
        - Provide a brief history of the organization and relevant context such as industry, region, size, and financial position.  
        - Include only details that are relevant for understanding the case.  
        - Ensure the tone is factual and objective.
        3. Focal Area and Internal Tension  
        - Describe the department, product, policy, or personnel involved.  
        - Highlight relevant organizational dynamics, cultural or political complexity, or competing interests.  
        - Include direct quotes or representative perspectives if available.
        4. Key Decision or Problem  
        - Present the dilemma or major decision the protagonist faces.  
        - Ensure the issue is open-ended, with no clear best answer.  
        - Avoid providing any form of analysis or hint at solutions.
        5. Alternatives  
        - Present two or three plausible options the protagonist is considering.  
        - Frame them neutrally as action choices.  
        - Avoid showing preference toward any option.
        6. Conclusion  
        - Conclude the case at the same time point introduced in the opening paragraph.  
        - Reiterate the challenge or uncertainty faced.  
        - End in a way that encourages classroom discussion.
        Writing Guidelines:
        - Write the entire case in past tense, except for direct quotes or exhibit captions.
        - Use plain, clear, and concise English that avoids complex sentence structures. Do not use idioms, metaphors, or regional expressions.
        - Maintain a neutral, journalistic tone throughout the case. Present facts and descriptions objectively, without showing personal bias or taking sides.
        - Do not include any analysis, interpretation, or recommendations. The case should not teach or explain the concepts — leave that for the classroom.
        - Avoid dramatic storytelling, flowery language, or fictionalized details (e.g., sunsets, internal thoughts, or emotionally loaded phrasing). Stick to the facts and actual events.
        - Keep paragraphs short and readable. Break information into logical sections that help the student understand the context and dilemma.
        - Ensure characters have clear roles and names. Introduce them by full name and position the first time they appear.
        - Use quotes only when necessary to reflect realistic perspectives or stakeholder voices. Do not overuse them.
        - Minimize use of abbreviations unless they are industry-standard. Always define an abbreviation the first time it's used.
        - Do not include the outcome of the decision or what happened afterward. The case must conclude at the same point in time as the opening paragraph.
        - Make sure the issue or decision point is realistic, debatable, and suitable for a 60–70 minute classroom discussion with ~2 hours of student preparation.

        Educational Objective:  
        The case should allow students to identify key issues, analyze multiple perspectives, evaluate competing trade-offs, and propose justifiable decisions. The case should be usable in a 60 to 70-minute classroom session and suitable for about 2 hours of student preparation time.
        """
        )

user_input = st.chat_input("Type your prompt or ask about the generated case study...")
if "last_case_study" in st.session_state and st.session_state.last_case_study:
    st.chat_message("assistant").markdown(st.session_state.last_case_study)
    docx_buffer = generate_docx_buffer(st.session_state.last_case_study)
    pdf_buffer = generate_pdf_buffer(st.session_state.last_case_study)

    col1, col2 = st.columns([1, 1])
    with col1:
        st.download_button("📄 Download .docx", docx_buffer, "case_study.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with col2:
        st.download_button("📝 Download .pdf", pdf_buffer, "case_study.pdf",
                           mime="application/pdf",
                           use_container_width=True)

if user_input or uploaded_files:
    if user_input:
        st.chat_message("user").markdown(user_input)
    elif uploaded_files:
        uploaded_names = ", ".join([f"`{file.name}`" for file in uploaded_files])
        st.chat_message("user").markdown(f"📂 Uploaded {uploaded_names}")

    if "last_case_study" not in st.session_state:
        all_text = ""
        for file in uploaded_files or []:
            all_text += extract_text_from_file(file)
            st.success(f"✅ {file.name} uploaded successfully.")
        content = all_text if all_text else user_input

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ]

        with st.spinner("Generating case study..."):
            try:
                response = client.chat.completions.create(
                    model='llama-3.3-70b-versatile',
                    messages=messages,
                    max_tokens=4096
                )
                case = response.choices[0].message.content
                st.chat_message("assistant").markdown(case)

                st.session_state.chat_history.append({"role": "user", "content": user_input})
                st.session_state.chat_history.append({"role": "assistant", "content": case})
                st.session_state.last_case_study = case

                docx_buffer = generate_docx_buffer(case)
                pdf_buffer = generate_pdf_buffer(case)

                col1, col2 = st.columns([1, 1])
                with col1:
                    st.download_button("📄 Download .docx", docx_buffer, "case_study.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True)
                with col2:
                    st.download_button("📝 Download .pdf", pdf_buffer, "case_study.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)

            except Exception as e:
                st.error(f"Error generating case study: {e}")

    else:
        followup_prompt = (
            "You are a highly interactive, intelligent, and emotionally aware case study assistant..."
        )

        followup_messages = [
            {
            "role": "system",
            "content": (
                "You are a highly interactive, intelligent, and emotionally aware case study assistant, designed to guide users through their exploration of a generated case study. "
                "You behave like a warm, helpful, and responsive human tutor — think of yourself as the perfect mix of a domain expert and a friendly chatbot like ChatGPT. "
                "You are capable of holding natural, flowing conversations and can interpret casual messages like greetings ('hi', 'hey'), gratitude ('thanks'), or general reactions ('cool', 'ok', 'interesting'). "
                "For such messages, reply in a light, friendly, and affirming tone — without forcing case-specific content unless it’s asked for. "

                "However, when a user asks a meaningful or analytical question related to the case study (like 'What’s the dilemma?', 'Who is the main actor?', or 'What are the options available?'), "
                "switch into expert mode. Provide thoughtful, clear, and informative answers **strictly based on the content of the case study**. Do not invent or hallucinate any information. "
                "If the answer cannot be supported by the given case, say so politely. You must not make assumptions or fabricate facts. "

                "Your tone should always be respectful, encouraging, and patient. Help users feel confident and supported in navigating the case. "
                "If the user’s question is vague or unclear, ask a clarifying question or offer gentle guidance on what they might want to explore. "
                "If a user becomes confused, reassure them and offer specific questions they might ask next. "

                "You may format responses with bullet points, line breaks, or clear structuring — but avoid sounding robotic. "
                "Your goal is to make the user feel like they’re chatting with a thoughtful, insightful person who truly understands both human conversation and the case content."

                "Never respond with 'I don't know.' Instead, say something like 'That detail isn’t included in the case, but here’s what we do know…'. "
                "If the user makes a mistake, correct them kindly and constructively. "

                "Above all, your mission is to support productive, engaging, and enjoyable discussion about the case — and make the user feel understood, heard, and helped."
            )
            },

            {"role": "user", "content": f"Case Study:\n{st.session_state.last_case_study}\n\nQuestion: {user_input}"}
        ]

        with st.spinner("Answering your question..."):
            try:
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=followup_messages,
                    max_tokens=1024
                )
                answer = response.choices[0].message.content
                st.chat_message("assistant").markdown(answer)
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
            except Exception as e:
                st.error(f"Error answering question: {e}")


