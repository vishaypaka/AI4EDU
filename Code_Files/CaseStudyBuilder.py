import os
import json
import streamlit as st
from groq import Groq
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import whisper
import tempfile
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import Paragraph
from docx import Document
from docx.shared import Inches

whisper_model = whisper.load_model("base")  # or "small", "medium", "large"

# Load API Key
working_dir = os.path.dirname(os.path.abspath(__file__))
config_data = json.load(open(f"{working_dir}/config.json"))
GROQ_API_KEY = config_data["GROQ_API_KEY"]
os.environ["GROQ_API_KEY"] = GROQ_API_KEY
client = Groq(api_key=GROQ_API_KEY)

st.set_page_config(page_title="Case Study Builder", page_icon="📚", layout="centered")
st.markdown("<h1 style='text-align: center;'>📚 Case Study Builder</h1>", unsafe_allow_html=True)

with st.expander("📂 Upload Case Study Documents", expanded=True):

    uploaded_files = st.file_uploader("Supported formats: PDF, Word, Text, Audio, or Video files", accept_multiple_files=True, type=['pdf', 'docx', 'txt', 'mp3', 'mp4', 'wav', 'm4a'])

if "last_case_study" not in st.session_state:
    st.session_state.last_case_study = ""
if "editable_case" not in st.session_state:
    st.session_state.editable_case = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

system_prompt = ("""
        You are a professional case study generator for a Program Manager leadership course.
        Use the uploaded content to generate a compelling case that follows internationally accepted academic case writing guidelines, combining AMRC, CAIS, and Penn State standards.
        Structure the output using the following six sections:
        1. Opening Paragraph  
        - Clearly state the date, location, and organization.  
        - Identify the main decision-maker by full name and role.  
        - Describe the event or issue trigger that brought the situation to light.  
        - Introduce the core problem or opportunity without revealing the outcome.  
        - Keep it concise and written in past tense, using a neutral and journalistic tone.
        2. Background  
        - Provide a brief history of the organization and relevant context such as industry, region, size, and financial position.  
        - Include only details that are relevant for understanding the case.  
        - Ensure the tone is factual and objective.
        3. Focal Area and Internal Tension  
        - Describe the department, product, policy, or personnel involved.  
        - Highlight relevant organizational dynamics, cultural or political complexity, or competing interests.  
        - Include direct quotes or representative perspectives if available.
        4. Key Decision or Problem  
        - Present the dilemma or major decision the protagonist faces.  
        - Ensure the issue is open-ended, with no clear best answer.  
        - Avoid providing any form of analysis or hint at solutions.
        5. Alternatives  
        - Present two or three plausible options the protagonist is considering.  
        - Frame them neutrally as action choices.  
        - Avoid showing preference toward any option.
        6. Conclusion  
        - Conclude the case at the same time point introduced in the opening paragraph.  
        - Reiterate the challenge or uncertainty faced.  
        - End in a way that encourages classroom discussion.
        Writing Guidelines:
        - Write the entire case in past tense, except for direct quotes or exhibit captions.
        - Use plain, clear, and concise English that avoids complex sentence structures. Do not use idioms, metaphors, or regional expressions.
        - Maintain a neutral, journalistic tone throughout the case. Present facts and descriptions objectively, without showing personal bias or taking sides.
        - Do not include any analysis, interpretation, or recommendations. The case should not teach or explain the concepts — leave that for the classroom.
        - Avoid dramatic storytelling, flowery language, or fictionalized details (e.g., sunsets, internal thoughts, or emotionally loaded phrasing). Stick to the facts and actual events.
        - Keep paragraphs short and readable. Break information into logical sections that help the student understand the context and dilemma.
        - Ensure characters have clear roles and names. Introduce them by full name and position the first time they appear.
        - Use quotes only when necessary to reflect realistic perspectives or stakeholder voices. Do not overuse them.
        - Minimize use of abbreviations unless they are industry-standard. Always define an abbreviation the first time it's used.
        - Do not include the outcome of the decision or what happened afterward. The case must conclude at the same point in time as the opening paragraph.
        - Make sure the issue or decision point is realistic, debatable, and suitable for a 60–70 minute classroom discussion with ~2 hours of student preparation.
        Educational Objective:  
        The case should allow students to identify key issues, analyze multiple perspectives, evaluate competing trade-offs, and propose justifiable decisions. The case should be usable in a 60 to 70-minute classroom session and suitable for about 2 hours of student preparation time.
""")

followup_prompt = ("""
You are a highly interactive assistant who helps revise and discuss case studies with users.
Whenever the user asks you to make updates to the generated case, use the latest editable version in context and return a revised version, maintaining the six-section structure and academic tone.
Do not hallucinate. Be realistic and stick to user instructions and existing text.
""")

def extract_text_from_file(file):
    text = ""
    try:
        if file.type == "application/pdf":
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_file = Document(file)
            for para in docx_file.paragraphs:
                text += para.text + "\n"
        elif file.type.startswith("text/"):
            text = file.read().decode("utf-8")
        elif file.type.startswith("audio/") or file.type.startswith("video/"):
            st.info(f"🔊 Transcribing {file.name} using Whisper...")
            text = transcribe_audio(file)
        else:
            st.warning(f"Unsupported file type: {file.name}")
    except Exception as e:
        st.warning(f"Couldn't read {file.name}: {e}")
    return text


def transcribe_audio(file):
    suffix = os.path.splitext(file.name)[1]  # Use correct file extension
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
        temp.write(file.read())
        temp.flush()
        temp_path = temp.name

    try:
        result = whisper_model.transcribe(temp_path)
        return result["text"]
    except Exception as e:
        st.warning(f"❌ Whisper transcription failed: {e}")
        return ""
    finally:
        os.remove(temp_path)  # Optional: clean up temp file

def parse_markdown_table(text):
    lines = [line.strip() for line in text.strip().split("\n") if "|" in line]
    if not lines or len(lines) < 2:
        return None
    return [row.strip('|').split('|') for row in lines if "---" not in row]

def generate_docx_buffer(content):
    buffer = BytesIO()
    doc = Document()
    sections = content.split("###")

    for section in sections:
        section = section.strip()
        if not section:
            continue

        lines = section.split("\n")
        title = lines[0].strip()
        body = "\n".join(lines[1:]).strip()

        # Add heading
        doc.add_heading(title, level=2)

        # Check if it's a table section
        if "|" in body and "---" in body:
            table_data = parse_markdown_table(body)
            if table_data:
                table = doc.add_table(rows=1, cols=len(table_data[0]))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, val in enumerate(table_data[0]):
                    hdr_cells[i].text = val.strip()

                for row in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = val.strip()
            else:
                doc.add_paragraph(body)
        else:
            for paragraph in body.split("\n\n"):
                doc.add_paragraph(paragraph.strip())

        doc.add_paragraph()  # spacer

    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_buffer(content):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors

    def parse_markdown_table(text):
        lines = [line.strip() for line in text.strip().split("\n") if "|" in line]
        if not lines or len(lines) < 2:
            return None
        return [row.strip('|').split('|') for row in lines if "---" not in row]

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

    sections = content.split("###")
    for section in sections:
        section = section.strip()
        if not section:
            continue
        lines = section.split("\n")
        title = lines[0].strip()
        body = "\n".join(lines[1:]).strip()

        flowables.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        flowables.append(Spacer(1, 6))

        if "|" in body and "---" in body:
            table_data_raw = parse_markdown_table(body)
            if table_data_raw:
                table_data = [
                    [Paragraph(cell.strip(), styles["BodyText"]) for cell in row]
                    for row in table_data_raw
                ]
                # Set fixed column widths to avoid overflow (adjust as needed)
                col_widths = [100] * len(table_data[0])  # or customize based on your columns
                table = Table(table_data, colWidths=col_widths, hAlign='LEFT', repeatRows=1)
    
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('WORDWRAP', (0, 0), (-1, -1), True),
                ]))
                flowables.append(table)
            else:
                flowables.append(Paragraph(body.replace("\n", "<br />"), styles["Normal"]))
        else:
            flowables.append(Paragraph(body.replace("\n", "<br />"), styles["Normal"]))

        flowables.append(Spacer(1, 12))

    doc.build(flowables)
    buffer.seek(0)
    return buffer



def generate_stakeholder_matrix(case_text):
    messages = [
        {"role": "system", "content": "You are an analyst extracting stakeholder information from case studies."},
        {"role": "user", "content": f"From the following case, create a stakeholder matrix table (Name | Role | Interest | Influence):\n\n{case_text}"}
    ]
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        max_tokens=512
    )
    return response.choices[0].message.content.strip()

def generate_stakeholder_perspectives(case_text):
    messages = [
        {"role": "system", "content": "You are a business analyst helping instructors understand stakeholder viewpoints."},
        {"role": "user", "content": f"From the case study below, create a table (Name | Position | Viewpoint on Issue):\n\n{case_text}"}
    ]
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        max_tokens=512
    )
    return response.choices[0].message.content.strip()

def generate_assignment_questions(case_text):
    messages = [
        {"role": "system", "content": "You are a university-level instructional designer."},
        {"role": "user", "content": f"Generate 3 reflective assignment questions for students based on this case:\n\n{case_text}"}
    ]
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        max_tokens=512
    )
    return response.choices[0].message.content.strip()

# Show full chat
for message in st.session_state.chat_history:
    st.chat_message(message["role"]).markdown(message["content"])

user_input = st.chat_input("Ask about the case, suggest changes, or say hello...")

if user_input:
    st.chat_message("user").markdown(user_input)
    st.session_state.chat_history.append({"role": "user", "content": user_input})

    if st.session_state.last_case_study:
        messages = [
            {"role": "system", "content": "You are a helpful AI tutor answering questions based on the following case study:\n\n" + st.session_state.last_case_study}
        ] + st.session_state.chat_history
    else:
        messages = st.session_state.chat_history

    response = client.chat.completions.create(
        model='llama-3.3-70b-versatile',
        messages=messages,
        max_tokens=1024
    )
    reply = response.choices[0].message.content
    st.chat_message("assistant").markdown(reply)
    st.session_state.chat_history.append({"role": "assistant", "content": reply})

# Handle file upload
if uploaded_files:
    all_text = ""
    for file in uploaded_files:
        all_text += extract_text_from_file(file)
        st.success(f"✅ {file.name} uploaded successfully.")

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": all_text}
    ]

    with st.spinner("Generating case study from uploaded content..."):
        try:
            response = client.chat.completions.create(
                model='llama-3.3-70b-versatile',
                messages=messages,
                max_tokens=4096
            )
            case = response.choices[0].message.content
            st.session_state.last_case_study = case
            st.session_state.editable_case = case

            matrix = generate_stakeholder_matrix(case)
            perspectives = generate_stakeholder_perspectives(case)
            questions = generate_assignment_questions(case)

            full_output = (
                f"{case}\n\n"
                f"### 🧾 Stakeholder Involvement Matrix\n{matrix}\n\n"
                f"### 🎯 Stakeholder Perspectives\n{perspectives}\n\n"
                f"### 📘 Assignment Questions\n{questions}"
            )

            st.chat_message("assistant").markdown(full_output)

            edited = st.text_area("✏ Edit Case and Add-ons", value=full_output, height=600)

            docx_buffer = generate_docx_buffer(edited)
            pdf_buffer = generate_pdf_buffer(edited)

            col1, col2 = st.columns(2)
            with col1:
                st.download_button("📄 Download DOCX", docx_buffer, "case_study.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            with col2:
                st.download_button("📝 Download PDF", pdf_buffer, "case_study.pdf", mime="application/pdf", use_container_width=True)

        except Exception as e:
            st.error(f"❌ Failed to generate: {e}")
